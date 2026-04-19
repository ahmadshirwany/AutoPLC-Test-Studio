from __future__ import annotations

import base64
from pathlib import Path
import re
import tempfile

import markdown
from docx import Document
from docx.shared import Inches
import requests

from .models import Artifact, GeneratedDocs


MAX_MERMAID_CODE_CHARS_FOR_DOCX = 14000
MAX_MERMAID_URL_CHARS = 7000


def _append_mermaid_section(markdown_text: str, heading: str, mermaid_code: str) -> str:
    if not mermaid_code.strip() or heading in markdown_text:
        return markdown_text

    return (
        markdown_text.rstrip()
        + f"\n\n## {heading}\n\n```mermaid\n{mermaid_code.strip()}\n```\n"
    )


def _inject_diagrams(kind: str, markdown_text: str, docs: GeneratedDocs, include_diagrams: bool) -> str:
    if not include_diagrams or not docs.diagrams:
        return markdown_text

    enriched = markdown_text
    system_flow = docs.diagrams.get("system_flow", "")
    logic_flow = docs.diagrams.get("logic_flow", "")

    if kind == "system_overview":
        enriched = _append_mermaid_section(enriched, "System Flow Diagram", system_flow)
        enriched = _append_mermaid_section(enriched, "Logic Flow Diagram", logic_flow)
    elif kind == "detailed_code_documentation":
        enriched = _append_mermaid_section(enriched, "Logic Flow Diagram", logic_flow)

    return enriched


def _markdown_to_html_document(markdown_text: str, title: str) -> str:
    body = markdown.markdown(markdown_text, extensions=["fenced_code", "tables"])
    uses_mermaid = "```mermaid" in markdown_text

    mermaid_script = ""
    if uses_mermaid:
        mermaid_script = """
  <script type=\"module\">
    import mermaid from \"https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs\";

    const codeBlocks = document.querySelectorAll(\"pre > code.language-mermaid\");
    for (const block of codeBlocks) {
      const wrapper = document.createElement(\"div\");
      wrapper.className = \"mermaid\";
      wrapper.textContent = block.textContent;
      block.parentElement.replaceWith(wrapper);
    }

    mermaid.initialize({ startOnLoad: true, securityLevel: \"loose\" });
    mermaid.run();
  </script>
"""

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{title}</title>
  <style>
    body {{ font-family: Georgia, serif; max-width: 900px; margin: 2rem auto; line-height: 1.65; color: #1f2937; padding: 0 1rem; }}
    h1, h2, h3 {{ color: #111827; }}
    code {{ background: #f3f4f6; padding: 0.1rem 0.3rem; border-radius: 4px; }}
    pre {{ background: #111827; color: #f9fafb; padding: 1rem; border-radius: 8px; overflow-x: auto; }}
    table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
    th, td {{ border: 1px solid #d1d5db; padding: 0.5rem; text-align: left; }}
    th {{ background: #f9fafb; }}
  </style>
</head>
<body>
{body}
{mermaid_script}
</body>
</html>
"""


def _fetch_mermaid_png(mermaid_code: str, warnings: list[str]) -> bytes | None:
    compact_code = re.sub(r"\n{2,}", "\n", mermaid_code.strip())
    compact_code = "\n".join(line.rstrip() for line in compact_code.splitlines() if line.strip())

    if len(compact_code) > MAX_MERMAID_CODE_CHARS_FOR_DOCX:
        warnings.append(
            "DOCX diagram image skipped: Mermaid graph is too large for external rendering. "
            "Markdown and HTML keep the full interactive diagram."
        )
        return None

    encoded = base64.urlsafe_b64encode(compact_code.encode("utf-8")).decode("ascii")
    diagram_url = f"https://mermaid.ink/img/{encoded}"

    if len(diagram_url) > MAX_MERMAID_URL_CHARS:
        warnings.append(
            "DOCX diagram image skipped: Mermaid graph exceeds URL length limits for the renderer."
        )
        return None

    try:
        response = requests.get(diagram_url, timeout=25)
        response.raise_for_status()
    except requests.HTTPError as exc:
        status_code = exc.response.status_code if exc.response is not None else "unknown"
        if status_code == 414:
            warnings.append(
                "DOCX diagram image skipped: renderer rejected the diagram as too large (HTTP 414)."
            )
        else:
            warnings.append(
                f"Unable to render Mermaid diagram image for DOCX (HTTP {status_code})."
            )
        return None
    except requests.RequestException as exc:
        warnings.append(
            f"Unable to render Mermaid diagram image for DOCX ({exc.__class__.__name__})."
        )
        return None

    if not response.content:
        warnings.append("Mermaid diagram renderer returned empty image content.")
        return None

    return response.content


def _insert_mermaid_image(document: Document, mermaid_code: str, warnings: list[str]) -> None:
    image_bytes = _fetch_mermaid_png(mermaid_code, warnings)
    if image_bytes is None:
        document.add_paragraph("[Diagram image unavailable. Mermaid source retained in markdown/html outputs.]")
        return

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_image:
        temp_image.write(image_bytes)
        temp_image_path = Path(temp_image.name)

    try:
        document.add_picture(str(temp_image_path), width=Inches(6.5))
    except Exception as exc:  # pragma: no cover - defensive in case of malformed image bytes
        warnings.append(f"Failed to embed Mermaid diagram image in DOCX: {exc}")
        document.add_paragraph("[Diagram image embedding failed for this section.]")
    finally:
        try:
            temp_image_path.unlink(missing_ok=True)
        except OSError:
            pass


def _markdown_to_docx(markdown_text: str, output_path: Path, warnings: list[str]) -> None:
    document = Document()
    in_code_block = False
    code_block_language = ""
    code_lines: list[str] = []

    for line in markdown_text.splitlines():
        stripped = line.rstrip()

        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_block_language = stripped[3:].strip().lower()
                code_lines = []
            else:
                if code_block_language == "mermaid":
                    _insert_mermaid_image(document, "\n".join(code_lines), warnings)
                else:
                    for code_line in code_lines:
                        document.add_paragraph(code_line)

                in_code_block = False
                code_block_language = ""
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(stripped)
            continue

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        document.add_paragraph(stripped)

    # Gracefully flush any unterminated code block.
    if in_code_block and code_lines:
        if code_block_language == "mermaid":
            _insert_mermaid_image(document, "\n".join(code_lines), warnings)
        else:
            for code_line in code_lines:
                document.add_paragraph(code_line)

    document.save(str(output_path))


def write_documents(
    output_dir: Path,
    docs: GeneratedDocs,
    formats: set[str],
    include_diagrams: bool = True,
) -> list[Artifact]:
    output_dir.mkdir(parents=True, exist_ok=True)

    artifacts: list[Artifact] = []
    render_warnings: list[str] = []
    file_map = {
        "system_overview": docs.system_overview_md,
        "detailed_code_documentation": docs.detailed_code_md,
    }

    for kind, raw_markdown_text in file_map.items():
        markdown_text = _inject_diagrams(kind, raw_markdown_text, docs, include_diagrams)

        if "markdown" in formats:
            file_name = f"{kind}.md"
            file_path = output_dir / file_name
            file_path.write_text(markdown_text, encoding="utf-8")
            artifacts.append(
                Artifact(
                    kind=kind,
                    format="markdown",
                    file_name=file_name,
                    relative_path=f"{output_dir.name}/{file_name}",
                    download_url=f"/output/{output_dir.name}/{file_name}",
                )
            )

        if "html" in formats:
            file_name = f"{kind}.html"
            file_path = output_dir / file_name
            file_path.write_text(
                _markdown_to_html_document(markdown_text, title=kind.replace("_", " ").title()),
                encoding="utf-8",
            )
            artifacts.append(
                Artifact(
                    kind=kind,
                    format="html",
                    file_name=file_name,
                    relative_path=f"{output_dir.name}/{file_name}",
                    download_url=f"/output/{output_dir.name}/{file_name}",
                )
            )

        if "docx" in formats:
            file_name = f"{kind}.docx"
            file_path = output_dir / file_name
            _markdown_to_docx(markdown_text, file_path, render_warnings)
            artifacts.append(
                Artifact(
                    kind=kind,
                    format="docx",
                    file_name=file_name,
                    relative_path=f"{output_dir.name}/{file_name}",
                    download_url=f"/output/{output_dir.name}/{file_name}",
                )
            )

    for warning in render_warnings:
        if warning not in docs.warnings:
            docs.warnings.append(warning)

    return artifacts
